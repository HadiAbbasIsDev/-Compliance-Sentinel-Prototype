"""
Script to extract text from ComplianceData .docx files,
chunk them, and store embeddings in Qdrant Cloud using
server-side inference (sentence-transformers/all-minilm-l6-v2).
"""

import os
import uuid
from docx import Document as DocxDocument
from qdrant_client import QdrantClient
from qdrant_client.http.models import (
    PointStruct,
    Distance,
    VectorParams,
    Document,
)

# ── Configuration ──────────────────────────────────────────────
QDRANT_URL = "https://4d49b355-5e45-47fb-a39e-53f5bf2db6f7.us-west-2-0.aws.cloud.qdrant.io:6333"
QDRANT_API_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJhY2Nlc3MiOiJtIn0.bOVdHaPArTpF0sSoFvHo7NqzLV0-lCQV2IPIYikqKB4"
COLLECTION_NAME = "compliance_docs"
EMBEDDING_MODEL = "sentence-transformers/all-minilm-l6-v2"
VECTOR_DIM = 384
DATA_DIR = os.path.join(os.path.dirname(__file__), "ComplianceData")

# Chunk settings (characters)
CHUNK_SIZE = 500       # ~roughly fits within 256-token context window
CHUNK_OVERLAP = 100


# ── Helper functions ───────────────────────────────────────────
def extract_text_from_docx(filepath: str) -> str:
    """Return all paragraph text from a .docx file."""
    doc = DocxDocument(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks by character count."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


# ── Main ───────────────────────────────────────────────────────
def main():
    # 1. Connect to Qdrant Cloud (with cloud_inference enabled)
    client = QdrantClient(
        url=QDRANT_URL,
        api_key=QDRANT_API_KEY,
        cloud_inference=True,
    )
    print("✅ Connected to Qdrant Cloud")
    print("   Existing collections:", client.get_collections())

    # 2. (Re)create the collection
    if client.collection_exists(COLLECTION_NAME):
        client.delete_collection(COLLECTION_NAME)
        print(f"🗑  Deleted existing collection '{COLLECTION_NAME}'")

    client.create_collection(
        collection_name=COLLECTION_NAME,
        vectors_config=VectorParams(
            size=VECTOR_DIM,
            distance=Distance.COSINE,
        ),
    )
    print(f"📦 Created collection '{COLLECTION_NAME}' (dim={VECTOR_DIM}, cosine)")

    # 3. Read all .docx files, chunk, and build points
    points: list[PointStruct] = []
    point_id = 1

    docx_files = sorted(
        f for f in os.listdir(DATA_DIR) if f.endswith(".docx")
    )
    print(f"\n📂 Found {len(docx_files)} .docx files in {DATA_DIR}")

    for filename in docx_files:
        filepath = os.path.join(DATA_DIR, filename)
        full_text = extract_text_from_docx(filepath)
        chunks = chunk_text(full_text)
        print(f"   📄 {filename}: {len(full_text)} chars → {len(chunks)} chunks")

        for idx, chunk in enumerate(chunks):
            points.append(
                PointStruct(
                    id=point_id,
                    payload={
                        "source_file": filename,
                        "chunk_index": idx,
                        "text": chunk,
                    },
                    vector=Document(
                        text=chunk,
                        model=EMBEDDING_MODEL,
                    ),
                )
            )
            point_id += 1

    print(f"\n🔢 Total points to upsert: {len(points)}")

    # 4. Upsert in batches of 64
    BATCH_SIZE = 64
    for i in range(0, len(points), BATCH_SIZE):
        batch = points[i : i + BATCH_SIZE]
        client.upsert(collection_name=COLLECTION_NAME, points=batch)
        print(f"   ⬆️  Upserted batch {i // BATCH_SIZE + 1} ({len(batch)} points)")

    # 5. Verify
    collection_info = client.get_collection(COLLECTION_NAME)
    print(f"\n✅ Done! Collection '{COLLECTION_NAME}' now has {collection_info.points_count} points.")

    # 6. Quick test query
    print("\n🔍 Running a test query: 'KYC customer due diligence'")
    results = client.query_points(
        collection_name=COLLECTION_NAME,
        query=Document(
            text="KYC customer due diligence",
            model=EMBEDDING_MODEL,
        ),
        limit=3,
    )
    for pt in results.points:
        print(f"   Score: {pt.score:.4f} | File: {pt.payload['source_file']} | Chunk #{pt.payload['chunk_index']}")
        print(f"   Text: {pt.payload['text'][:120]}...\n")


if __name__ == "__main__":
    main()
